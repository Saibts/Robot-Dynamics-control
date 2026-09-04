"""
Script to generate the comprehensive, step-by-step execution guide PDF and Word document
specifically tailored for ROS 2 Jazzy Jalisco on Ubuntu 24.04 LTS.
"""
import os
import sys
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable, KeepTogether
)
from reportlab.pdfgen import canvas
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class NumberedGuideCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        if self._pageNumber == 1:
            return
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#4A5568"))

        # Header
        self.drawString(54, 750, "ROS 2 Jazzy (Ubuntu 24.04) — Step-by-Step Implementation & Execution Guide")
        self.drawRightString(612 - 54, 750, "Team 3 | RDC Project")
        self.setStrokeColor(colors.HexColor("#CBD5E0"))
        self.setLineWidth(0.75)
        self.line(54, 742, 612 - 54, 742)

        # Footer
        self.line(54, 45, 612 - 54, 45)
        self.drawString(54, 34, "CONFIDENTIAL TUTORIAL & LAB EXECUTION MANUAL — TEAM 3")
        page_str = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(612 - 54, 34, page_str)
        self.restoreState()

def build_step_by_step_pdf(filename="ROS2_Jazzy_Step_By_Step_Execution_Guide.pdf"):
    doc = SimpleDocTemplate(
        filename,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()

    navy = colors.HexColor("#1A365D")
    slate = colors.HexColor("#2B6CB0")
    dark = colors.HexColor("#2D3748")
    bg_box = colors.HexColor("#F7FAFC")
    code_bg = colors.HexColor("#1A202C")
    code_text = colors.HexColor("#68D391")

    p_cover_title = ParagraphStyle('CTitle', fontName='Helvetica-Bold', fontSize=22, leading=26, textColor=navy, alignment=1, spaceAfter=10)
    p_cover_sub = ParagraphStyle('CSub', fontName='Helvetica', fontSize=11, leading=15, textColor=slate, alignment=1, spaceAfter=20)
    p_h1 = ParagraphStyle('H1', fontName='Helvetica-Bold', fontSize=13, leading=16, textColor=navy, spaceBefore=12, spaceAfter=6, keepWithNext=True)
    p_h2 = ParagraphStyle('H2', fontName='Helvetica-Bold', fontSize=10.5, leading=13, textColor=slate, spaceBefore=8, spaceAfter=3, keepWithNext=True)
    p_body = ParagraphStyle('Body', fontName='Helvetica', fontSize=8.5, leading=12, textColor=dark, spaceAfter=4)
    p_bullet = ParagraphStyle('Bullet', fontName='Helvetica', fontSize=8.5, leading=12, textColor=dark, leftIndent=12, firstLineIndent=-8, spaceAfter=3)

    story = []

    # Cover Page
    story.append(Spacer(1, 30))
    story.append(Paragraph("<b>DEPARTMENT OF ROBOTICS ENGINEERING</b>", ParagraphStyle('D', fontName='Helvetica-Bold', fontSize=11, leading=13, textColor=slate, alignment=1, spaceAfter=6)))
    story.append(Paragraph("COURSE: ROBOTICS, DYNAMICS & CONTROL (SEM 5 / III YEAR)", ParagraphStyle('C', fontName='Helvetica', fontSize=9.5, leading=12, textColor=colors.HexColor("#718096"), alignment=1, spaceAfter=20)))
    story.append(HRFlowable(width="80%", thickness=2, color=navy, spaceAfter=15, spaceBefore=0))
    story.append(Paragraph("STEP-BY-STEP IMPLEMENTATION & SIMULATION EXECUTION MANUAL", p_cover_title))
    story.append(Paragraph("Heterogeneous Multi-Robot Coordination (UR5 + TurtleBot3 OpenManipulator) Using ROS 2 Actions & Services in ROS 2 Jazzy Jalisco on Ubuntu 24.04 LTS", p_cover_sub))
    story.append(HRFlowable(width="80%", thickness=2, color=navy, spaceAfter=25, spaceBefore=0))

    # Meta table
    meta_data = [
        [Paragraph("<b>Target Operating System:</b>", ParagraphStyle('M1', fontName='Helvetica-Bold', fontSize=8.5, textColor=navy)), Paragraph("Ubuntu 24.04 LTS (Noble Numbat)", ParagraphStyle('M2', fontName='Helvetica', fontSize=8.5, textColor=dark))],
        [Paragraph("<b>ROS 2 Distribution:</b>", ParagraphStyle('M1', fontName='Helvetica-Bold', fontSize=8.5, textColor=navy)), Paragraph("<b>ROS 2 Jazzy Jalisco</b> (Latest Tier 1 LTS Release)", ParagraphStyle('M2', fontName='Helvetica', fontSize=8.5, textColor=dark))],
        [Paragraph("<b>Assigned Project:</b>", ParagraphStyle('M1', fontName='Helvetica-Bold', fontSize=8.5, textColor=navy)), Paragraph("Team 3: Multi-Robot Coordination Using ROS 2 Actions & Services", ParagraphStyle('M2', fontName='Helvetica', fontSize=8.5, textColor=dark))],
        [Paragraph("<b>Designated Robots:</b>", ParagraphStyle('M1', fontName='Helvetica-Bold', fontSize=8.5, textColor=navy)), Paragraph("1) Fixed Manipulator: 6-DOF Universal Robot (UR5)<br/>2) Mobile Manipulator: TurtleBot3 (Waffle Pi) + OpenManipulator-X", ParagraphStyle('M2', fontName='Helvetica', fontSize=8.5, textColor=dark))],
        [Paragraph("<b>Simulation Platforms:</b>", ParagraphStyle('M1', fontName='Helvetica-Bold', fontSize=8.5, textColor=navy)), Paragraph("Gazebo Harmonic (ros_gz), RViz2, Nav2, MoveIt 2, ros2_control", ParagraphStyle('M2', fontName='Helvetica', fontSize=8.5, textColor=dark))]
    ]
    t_meta = Table(meta_data, colWidths=[140, 330])
    t_meta.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), bg_box),
        ('BOX', (0,0), (-1,-1), 1.2, navy),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E0")),
        ('PADDING', (0,0), (-1,-1), 5),
    ]))
    story.append(t_meta)
    story.append(Spacer(1, 20))

    # Note
    note_box = [
        [Paragraph("<b>ABOUT THIS PRACTICAL GUIDE:</b><br/>"
                   "This guide is written specifically for students starting from scratch on Ubuntu Linux with ROS 2 Jazzy. Every command is presented in copy-paste ready format, explaining what each terminal does, how the simulation behaves, and how to verify that your ROS 2 Actions and Services are working properly.", ParagraphStyle('NB', fontName='Helvetica', fontSize=8.5, leading=12, textColor=colors.HexColor("#2C5282")))]
    ]
    t_note = Table(note_box, colWidths=[470])
    t_note.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#EBF8FF")),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor("#BEE3F8")),
        ('PADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(t_note)

    story.append(PageBreak())

    def add_code_block(code_str):
        code_formatted = code_str.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>').replace(' ', '&nbsp;')
        p_c = Paragraph(code_formatted, ParagraphStyle('CB', fontName='Courier', fontSize=7.5, leading=10, textColor=code_text))
        t_c = Table([[p_c]], colWidths=[490])
        t_c.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), code_bg),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor("#2D3748")),
            ('PADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(t_c)
        story.append(Spacer(1, 4))

    # ==========================================
    # PHASE 1: ENVIRONMENT & PACKAGES
    # ==========================================
    story.append(Paragraph("Phase 1: Ubuntu 24.04 & ROS 2 Jazzy Prerequisites Setup", p_h1))
    story.append(Paragraph(
        "Before building the multi-robot project, ensure your Ubuntu 24.04 environment has all necessary ROS 2 Jazzy navigation, manipulation, and simulation libraries installed.",
        p_body
    ))

    story.append(Paragraph("Step 1.1: Verify ROS 2 Jazzy Installation", p_h2))
    story.append(Paragraph("Open a terminal (<code>Ctrl + Alt + T</code>) and verify that ROS 2 Jazzy is active:", p_body))
    add_code_block(
        "source /opt/ros/jazzy/setup.bash\n"
        "echo \"Active ROS Distro: $ROS_DISTRO\"\n"
        "# Expected Output: Active ROS Distro: jazzy"
    )

    story.append(Paragraph("Step 1.2: Install Required ROS 2 Jazzy Packages", p_h2))
    story.append(Paragraph("Run the following commands to install Nav2, MoveIt 2, Gazebo bridge, and robot packages:", p_body))
    add_code_block(
        "sudo apt update\n"
        "sudo apt install -y \\\n"
        "  ros-jazzy-navigation2 \\\n"
        "  ros-jazzy-nav2-bringup \\\n"
        "  ros-jazzy-moveit \\\n"
        "  ros-jazzy-moveit-ros-planning-interface \\\n"
        "  ros-jazzy-ros2-control \\\n"
        "  ros-jazzy-ros2-controllers \\\n"
        "  ros-jazzy-ros-gz \\\n"
        "  ros-jazzy-turtlebot3 \\\n"
        "  ros-jazzy-turtlebot3-msgs \\\n"
        "  ros-jazzy-turtlebot3-simulations \\\n"
        "  ros-jazzy-dynamixel-sdk \\\n"
        "  python3-colcon-common-extensions \\\n"
        "  python3-rosdep \\\n"
        "  python3-pip"
    )

    story.append(Paragraph("Step 1.3: Initialize and Update Rosdep", p_h2))
    add_code_block(
        "sudo rosdep init 2>/dev/null || true\n"
        "rosdep update"
    )

    story.append(Spacer(1, 8))

    # ==========================================
    # PHASE 2: WORKSPACE SETUP & CODE INTEGRATION
    # ==========================================
    story.append(Paragraph("Phase 2: Setting Up the Multi-Robot Colcon Workspace", p_h1))
    story.append(Paragraph(
        "Follow these steps to create your ROS 2 workspace and configure the <code>multi_robot_coordination</code> package.",
        p_body
    ))

    story.append(Paragraph("Step 2.1: Create Workspace Directory Structure", p_h2))
    add_code_block(
        "mkdir -p ~/ros2_ws/src/multi_robot_coordination/action\n"
        "mkdir -p ~/ros2_ws/src/multi_robot_coordination/srv\n"
        "mkdir -p ~/ros2_ws/src/multi_robot_coordination/multi_robot_coordination\n"
        "mkdir -p ~/ros2_ws/src/multi_robot_coordination/launch\n"
        "mkdir -p ~/ros2_ws/src/multi_robot_coordination/config\n"
        "cd ~/ros2_ws/src/multi_robot_coordination"
    )

    story.append(Paragraph("Step 2.2: Copy Project Files from Workspace to Ubuntu", p_h2))
    story.append(Paragraph(
        "Copy all provided Python nodes, action interfaces, service interfaces, launch files, and package manifests from the project repository into <code>~/ros2_ws/src/multi_robot_coordination/</code>.",
        p_body
    ))
    story.append(Paragraph(
        "Directory check: verify that your folder layout matches:",
        p_body
    ))
    add_code_block(
        "~/ros2_ws/src/multi_robot_coordination/\n"
        "├── action/\n"
        "│   ├── NavigateAndPick.action\n"
        "│   └── UR5PickAndPlace.action\n"
        "├── srv/\n"
        "│   ├── AcquireHandoffLock.srv\n"
        "│   └── TriggerGripper.srv\n"
        "├── multi_robot_coordination/\n"
        "│   ├── __init__.py\n"
        "│   ├── scheduler.py\n"
        "│   ├── simulation_runner.py\n"
        "│   ├── multi_robot_coordinator.py\n"
        "│   ├── tb3_action_server.py\n"
        "│   ├── ur5_action_server.py\n"
        "│   └── handoff_service_server.py\n"
        "├── launch/\n"
        "│   └── multi_robot_system.launch.py\n"
        "├── package.xml\n"
        "└── setup.py"
    )

    story.append(PageBreak())

    # ==========================================
    # PHASE 3: COMPILATION & SOURCING
    # ==========================================
    story.append(Paragraph("Phase 3: Building and Sourcing the Package", p_h1))
    story.append(Paragraph(
        "Compile the workspace using <code>colcon</code> and source the overlay environment.",
        p_body
    ))

    story.append(Paragraph("Step 3.1: Build Workspace with Colcon", p_h2))
    add_code_block(
        "cd ~/ros2_ws\n"
        "rosdep install --from-paths src --ignore-src -r -y\n"
        "colcon build --symlink-install"
    )

    story.append(Paragraph("Step 3.2: Source the Built Workspace Overlay", p_h2))
    add_code_block(
        "source install/setup.bash\n"
        "# To make sourcing permanent in your terminal:\n"
        "echo \"source ~/ros2_ws/install/setup.bash\" >> ~/.bashrc"
    )

    story.append(Paragraph("Step 3.3: Verify Installed Interfaces", p_h2))
    add_code_block(
        "ros2 interface show multi_robot_coordination/action/NavigateAndPick\n"
        "ros2 interface show multi_robot_coordination/srv/AcquireHandoffLock"
    )

    story.append(Spacer(1, 8))

    # ==========================================
    # PHASE 4: RUNNING THE MULTI-ROBOT SYSTEM
    # ==========================================
    story.append(Paragraph("Phase 4: Step-by-Step Execution Guide (Terminal by Terminal)", p_h1))
    story.append(Paragraph(
        "To run the complete coordinated multi-robot simulation, open multiple terminal windows or use <code>tmux</code> / <code>terminator</code> tabs. Follow the exact sequence below:",
        p_body
    ))

    terminals = [
        ("TERMINAL 1: Shared Transfer Zone Lock Service Server",
         "Launches the mutual exclusion lock server that manages exclusive access to the transfer zone:\n"
         "source ~/ros2_ws/install/setup.bash\n"
         "ros2 run multi_robot_coordination lock_server\n"
         "# Expected: [INFO] [handoff_lock_server]: Shared Transfer Zone Lock Service Server Active."),

        ("TERMINAL 2: TurtleBot3 Mobile Manipulator Action Server",
         "Launches the Action Server simulating mobile base navigation and OpenManipulator-X grasping:\n"
         "source ~/ros2_ws/install/setup.bash\n"
         "ros2 run multi_robot_coordination tb3_server\n"
         "# Expected: [INFO] [tb3_action_server]: TurtleBot3 Mobile Manipulator Action Server Started."),

        ("TERMINAL 3: UR5 Fixed Manipulator Action Server",
         "Launches the Action Server executing 6-DOF MoveIt 2 trajectory planning and assembly:\n"
         "source ~/ros2_ws/install/setup.bash\n"
         "ros2 run multi_robot_coordination ur5_server\n"
         "# Expected: [INFO] [ur5_action_server]: UR5 6-DOF Manipulator Action Server Started."),

        ("TERMINAL 4: Central Multi-Robot Coordinator Node",
         "Launches the main orchestrator with Priority-Based scheduling:\n"
         "source ~/ros2_ws/install/setup.bash\n"
         "ros2 run multi_robot_coordination coordinator --ros-args -p scheduler_mode:=PRIORITY\n"
         "# Expected: [INFO] [multi_robot_coordinator]: Task Scheduler Mode set to: PRIORITY"),

        ("TERMINAL 5: (All-in-One Alternative) Launching Entire System via Launch File",
         "Instead of opening 4 individual terminals, launch everything at once using the master launch file:\n"
         "source ~/ros2_ws/install/setup.bash\n"
         "ros2 launch multi_robot_coordination multi_robot_system.launch.py")
    ]

    for t_name, t_cmd in terminals:
        story.append(Paragraph(f"<b>{t_name}</b>", p_h2))
        add_code_block(t_cmd)

    story.append(PageBreak())

    # ==========================================
    # PHASE 5: INTERACTIVE TESTING & MONITORING
    # ==========================================
    story.append(Paragraph("Phase 5: Interactive Testing, Introspection & Monitoring", p_h1))
    story.append(Paragraph(
        "Open a new terminal to inspect the active ROS 2 Actions, Services, Topics, and test triggering operations manually:",
        p_body
    ))

    story.append(Paragraph("5.1: Inspect Active ROS 2 Actions & Services", p_h2))
    add_code_block(
        "# List all active Action Servers\n"
        "ros2 action list\n"
        "# Output: /navigate_and_pick\n"
        "#         /ur5_pick_and_assemble\n\n"
        "# List all active Services\n"
        "ros2 service list | grep -E 'lock|gripper'\n"
        "# Output: /acquire_transfer_lock"
    )

    story.append(Paragraph("5.2: Send a Manual Test Goal to TurtleBot3 Action Server", p_h2))
    add_code_block(
        "ros2 action send_goal --feedback /navigate_and_pick multi_robot_coordination/action/NavigateAndPick \\\n"
        "  \"{target_station_id: 'STATION_A', pickup_coordinates: [1.2, 0.8, 0.15], priority_level: 1}\"\n"
        "# You will see real-time 10 Hz feedback streamed to your terminal until completion!"
    )

    story.append(Paragraph("5.3: Test Shared Zone Lock Service", p_h2))
    add_code_block(
        "# Request lock acquisition:\n"
        "ros2 service call /acquire_transfer_lock multi_robot_coordination/srv/AcquireHandoffLock \\\n"
        "  \"{robot_id: 'turtlebot3', zone_id: 1, request_lock: true}\"\n\n"
        "# Release lock:\n"
        "ros2 service call /acquire_transfer_lock multi_robot_coordination/srv/AcquireHandoffLock \\\n"
        "  \"{robot_id: 'turtlebot3', zone_id: 1, request_lock: false}\""
    )

    story.append(Spacer(1, 8))

    # ==========================================
    # PHASE 6: RUNNING THE BENCHMARK EXPERIMENT
    # ==========================================
    story.append(Paragraph("Phase 6: Running the Automated Benchmark Suite & Generating Results", p_h1))
    story.append(Paragraph(
        "To evaluate and reproduce the FIFO, Priority-Based, and Round-Robin benchmark metrics reported in the project submission:",
        p_body
    ))
    add_code_block(
        "cd ~/ros2_ws/src/multi_robot_coordination/multi_robot_coordination\n"
        "python3 simulation_runner.py"
    )
    story.append(Paragraph(
        "<b>Expected Terminal Output:</b><br/>"
        "The script will run 30 stochastic assembly jobs across all 3 algorithms and print the exact metrics table (Makespan, Avg Wait Time, Priority-1 Wait Time, UR5 Utilization %, TB3 Utilization %, and Tasks/Hour).",
        p_body
    ))

    story.append(Spacer(1, 8))

    # ==========================================
    # PHASE 7: JAZZY TROUBLESHOOTING TIPS
    # ==========================================
    story.append(Paragraph("Phase 7: ROS 2 Jazzy Specific Troubleshooting Tips", p_h1))
    
    jazzy_tips = [
        ("Issue: <code>colcon build</code> fails with <code>rosidl_default_generators not found</code>",
         "Fix: Run <code>sudo apt install ros-jazzy-rosidl-default-generators ros-jazzy-rosidl-default-runtime</code> and re-run <code>colcon build</code>."),
        
        ("Issue: <code>ros2: command not found</code> or action interfaces cannot be imported in Python",
         "Fix: Make sure you ran <code>source /opt/ros/jazzy/setup.bash</code> followed by <code>source ~/ros2_ws/install/setup.bash</code> in EVERY open terminal."),
        
        ("Issue: Action Server reports <code>wait_for_server()</code> timeout",
         "Fix: Check if the nodes are running under different ROS_DOMAIN_ID. In each terminal run: <code>export ROS_DOMAIN_ID=42</code> and restart the nodes."),
        
        ("Issue: Permission denied when running Python node scripts",
         "Fix: Give executable permissions to all Python files: <code>chmod +x ~/ros2_ws/src/multi_robot_coordination/multi_robot_coordination/*.py</code>.")
    ]

    for issue, fix in jazzy_tips:
        story.append(Paragraph(f"<b>{issue}</b>", ParagraphStyle('JTI', fontName='Helvetica-Bold', fontSize=8.5, leading=11, textColor=colors.HexColor("#C53030"))))
        story.append(Paragraph(fix, ParagraphStyle('JTF', fontName='Helvetica', fontSize=8, leading=11, textColor=dark, spaceAfter=4)))

    doc.build(story, canvasmaker=NumberedGuideCanvas)
    print(f"Step-by-step PDF generated successfully: {filename}")

def build_step_by_step_docx(filename="ROS2_Jazzy_Step_By_Step_Execution_Guide.docx"):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    navy = RGBColor(26, 54, 93)
    slate = RGBColor(43, 108, 176)
    dark = RGBColor(45, 55, 72)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run("STEP-BY-STEP IMPLEMENTATION & SIMULATION EXECUTION GUIDE\nROS 2 JAZZY JALISCO (UBUNTU 24.04 LTS)")
    r_t.font.name = "Calibri"
    r_t.font.size = Pt(16)
    r_t.font.bold = True
    r_t.font.color.rgb = navy

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s = p_sub.add_run("Team 3: Multi-Robot Coordination Using ROS 2 Actions & Services (UR5 + TurtleBot3)")
    r_s.font.name = "Calibri"
    r_s.font.size = Pt(11)
    r_s.font.italic = True
    r_s.font.color.rgb = slate

    doc.add_paragraph()

    sections = [
        ("Phase 1: Environment & Prerequisites (Ubuntu 24.04 + ROS 2 Jazzy)",
         "1. Verify ROS 2 Jazzy:\nsource /opt/ros/jazzy/setup.bash\necho $ROS_DISTRO\n\n"
         "2. Install Dependencies:\nsudo apt update && sudo apt install -y ros-jazzy-navigation2 ros-jazzy-nav2-bringup ros-jazzy-moveit ros-jazzy-moveit-ros-planning-interface ros-jazzy-ros2-control ros-jazzy-ros2-controllers ros-jazzy-ros-gz ros-jazzy-turtlebot3 ros-jazzy-turtlebot3-msgs python3-colcon-common-extensions python3-rosdep python3-pip\n\n"
         "3. Update Rosdep:\nsudo rosdep init 2>/dev/null || true && rosdep update"),

        ("Phase 2: Workspace Setup & Code Layout",
         "1. Create Directories:\nmkdir -p ~/ros2_ws/src/multi_robot_coordination/{action,srv,multi_robot_coordination,launch,config}\n\n"
         "2. Place all provided files into ~/ros2_ws/src/multi_robot_coordination/"),

        ("Phase 3: Compilation with Colcon",
         "cd ~/ros2_ws\n"
         "rosdep install --from-paths src --ignore-src -r -y\n"
         "colcon build --symlink-install\n"
         "source install/setup.bash"),

        ("Phase 4: Terminal-by-Terminal Execution Sequence",
         "Terminal 1 (Lock Server):\nros2 run multi_robot_coordination lock_server\n\n"
         "Terminal 2 (TurtleBot3 Server):\nros2 run multi_robot_coordination tb3_server\n\n"
         "Terminal 3 (UR5 Server):\nros2 run multi_robot_coordination ur5_server\n\n"
         "Terminal 4 (Coordinator Node):\nros2 run multi_robot_coordination coordinator --ros-args -p scheduler_mode:=PRIORITY\n\n"
         "OR All-In-One Master Launch:\nros2 launch multi_robot_coordination multi_robot_system.launch.py"),

        ("Phase 5: Interactive Introspection & Benchmark Execution",
         "1. Check Actions & Services:\nros2 action list\nros2 service list | grep lock\n\n"
         "2. Run Automated Benchmark Suite:\npython3 ~/ros2_ws/src/multi_robot_coordination/multi_robot_coordination/simulation_runner.py")
    ]

    for s_title, s_content in sections:
        h = doc.add_heading(s_title, level=2)
        h.style.font.color.rgb = navy
        p = doc.add_paragraph(s_content)
        p.style.font.color.rgb = dark

    doc.save(filename)
    print(f"Step-by-step Word doc saved: {filename}")

if __name__ == '__main__':
    build_step_by_step_pdf()
    build_step_by_step_docx()
